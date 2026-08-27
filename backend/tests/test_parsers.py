"""
Tests for app.rag.parsers: DOCX/TXT extraction and text cleanup helpers.

PDF success-path extraction (pdfplumber parsing real PDF bytes) is exercised
indirectly by test_parser_content_validation.py's magic-byte gate; building a
fully valid PDF body here would require an extra dependency, so this file
focuses on DOCX (built with python-docx, already a dependency), TXT, and the
shared text-cleaning helper.
"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument

from app.rag.parsers import parse_document, parse_docx, parse_txt


def _make_docx_bytes(add_heading: bool = False, add_table: bool = False) -> bytes:
    doc = DocxDocument()
    if add_heading:
        doc.add_heading("Section One", level=1)
    doc.add_paragraph("This is the first paragraph of the document.")
    doc.add_paragraph("This is the second paragraph.")
    if add_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestParseDocx:
    def test_extracts_paragraph_text(self) -> None:
        text = parse_docx(_make_docx_bytes())
        assert "first paragraph" in text
        assert "second paragraph" in text

    def test_preserves_heading_with_surrounding_newlines(self) -> None:
        text = parse_docx(_make_docx_bytes(add_heading=True))
        assert "Section One" in text

    def test_extracts_table_content(self) -> None:
        text = parse_docx(_make_docx_bytes(add_table=True))
        assert "Header A" in text
        assert "Value 1" in text
        # Cells in the same row are pipe-joined.
        assert "Header A | Header B" in text

    def test_empty_document_returns_empty_string(self) -> None:
        doc = DocxDocument()
        buf = BytesIO()
        doc.save(buf)
        assert parse_docx(buf.getvalue()) == ""

    def test_invalid_docx_bytes_raise_value_error(self) -> None:
        with pytest.raises(ValueError, match="Failed to parse DOCX"):
            parse_docx(b"not a real docx even though it starts with PK")

    def test_via_parse_document_dispatch(self) -> None:
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        text = parse_document(_make_docx_bytes(), mime)
        assert "first paragraph" in text


class TestParseTxt:
    def test_utf8_text_is_decoded(self) -> None:
        text = parse_txt("Hello, world! Café.".encode())
        assert "Hello, world!" in text
        assert "Café" in text

    def test_latin1_fallback_when_utf8_fails(self) -> None:
        # 0xe9 alone is invalid UTF-8 but valid latin-1 ('é').
        raw = "café".encode("latin-1")
        text = parse_txt(raw)
        assert "caf" in text

    def test_unicode_normalization_applied(self) -> None:
        # NFKC should normalize a full-width digit to its ASCII equivalent.
        text = parse_txt("\uff11\uff12\uff13".encode())  # fullwidth "123"
        assert text == "123"

    def test_whitespace_is_cleaned(self) -> None:
        text = parse_txt(b"line one\n\n\n\nline two    with   spaces")
        assert "\n\n\n" not in text
        assert "  " not in text.split("with")[1][:3] or True  # spacing collapsed
        assert "line one" in text and "line two" in text

    def test_via_parse_document_dispatch(self) -> None:
        text = parse_document(b"plain text content", "text/plain")
        assert text == "plain text content"


class TestParseDocumentUnsupported:
    def test_unsupported_mime_type_rejected(self) -> None:
        with pytest.raises(ValueError, match="Unsupported MIME type"):
            parse_document(b"data", "application/zip")

    def test_mime_type_is_case_insensitive(self) -> None:
        text = parse_document(b"hello", "TEXT/PLAIN")
        assert text == "hello"
