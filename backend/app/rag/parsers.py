"""
Document parsers for extracting text from various file formats.
"""
import re
import unicodedata
from io import BytesIO

import pdfplumber
import structlog
from docx import Document as DocxDocument

logger = structlog.get_logger(__name__)


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using pdfplumber.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Cleaned text string
        
    Raises:
        ValueError: If PDF parsing fails
    """
    try:
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    
        if not text_parts:
            logger.warning("pdf_no_text_extracted")
            return ""
            
        full_text = "\n\n".join(text_parts)
        
        # Clean up the text
        full_text = _clean_text(full_text)
        
        # Remove page numbers pattern (Page X of Y, Page X, etc.)
        full_text = re.sub(r'\bPage\s+\d+\s*(of\s+\d+)?\b', '', full_text, flags=re.IGNORECASE)
        
        logger.info("pdf_parsed_successfully", pages=len(text_parts))
        return full_text
        
    except Exception as exc:
        logger.error("pdf_parsing_failed", error=str(exc))
        raise ValueError(f"Failed to parse PDF: {str(exc)}") from exc


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using python-docx.
    Preserves heading hierarchy and table content.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Cleaned text string
        
    Raises:
        ValueError: If DOCX parsing fails
    """
    try:
        doc = DocxDocument(BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs with heading preservation
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Add extra newline before headings for hierarchy
                if paragraph.style.name.startswith('Heading'):
                    text_parts.append(f"\n{text}\n")
                else:
                    text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_text = _extract_table_text(table)
            if table_text:
                text_parts.append(f"\n{table_text}\n")
        
        if not text_parts:
            logger.warning("docx_no_text_extracted")
            return ""
            
        full_text = "\n".join(text_parts)
        full_text = _clean_text(full_text)
        
        logger.info("docx_parsed_successfully", paragraphs=len(doc.paragraphs), tables=len(doc.tables))
        return full_text
        
    except Exception as exc:
        logger.error("docx_parsing_failed", error=str(exc))
        raise ValueError(f"Failed to parse DOCX: {str(exc)}") from exc


def _extract_table_text(table) -> str:
    """Extract text from a DOCX table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):  # Only include non-empty rows
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def parse_txt(file_bytes: bytes) -> str:
    """
    Extract text from plain text file.
    Handles encoding and unicode normalization.
    
    Args:
        file_bytes: Text file content as bytes
        
    Returns:
        Cleaned text string
        
    Raises:
        ValueError: If text decoding fails
    """
    try:
        # Try UTF-8 first, fall back to latin-1
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        
        # Normalize unicode (NFKC: Compatibility decomposition followed by canonical composition)
        text = unicodedata.normalize('NFKC', text)
        
        # Clean and strip
        text = _clean_text(text)
        
        logger.info("txt_parsed_successfully", length=len(text))
        return text
        
    except Exception as exc:
        logger.error("txt_parsing_failed", error=str(exc))
        raise ValueError(f"Failed to parse text file: {str(exc)}") from exc


def parse_document(file_bytes: bytes, mime_type: str) -> str:
    """
    Route to appropriate parser based on MIME type.
    
    Args:
        file_bytes: File content as bytes
        mime_type: MIME type of the file
        
    Returns:
        Extracted text string
        
    Raises:
        ValueError: If mime_type is unsupported or parsing fails
    """
    mime_type = mime_type.lower()
    
    if mime_type == "application/pdf":
        return parse_pdf(file_bytes)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(file_bytes)
    elif mime_type == "text/plain":
        return parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported MIME type: {mime_type}")


def _clean_text(text: str) -> str:
    """
    Clean up extracted text by removing excessive whitespace.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Replace multiple newlines with double newline
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text
